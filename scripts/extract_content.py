#!/usr/bin/env python3
"""
Multi-format Content Extractor for knowledge2skills

Supports: .pdf, .md, .txt, .docx, .epub, .mobi, .azw3, .json, .png, .jpg, .jpeg, .webp
Can process a single file or a list of files.
"""

import os
import json
import re
import argparse
import sys
import subprocess
import base64
import mimetypes
from pathlib import Path

# Optional dependency imports
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
except ImportError:
    epub = None

try:
    import mobi
except ImportError:
    mobi = None

try:
    import requests
except ImportError:
    requests = None

try:
    from PIL import Image
except ImportError:
    Image = None

# MinerU (Magic-PDF) support
MAGIC_PDF_AVAILABLE = False
LOCAL_MINERU_PYTHON = os.path.expanduser("~/.pyenv/versions/3.11.9/bin/python")
MINERU_LOCAL_TIMEOUT_SECONDS = int(os.environ.get("KNOWLEDGE2SKILLS_MINERU_TIMEOUT", "1800"))

try:
    from magic_pdf.pipe.UNIPipe import UNIPipe
    from magic_pdf.rw.AbsReaderWriter import FileReaderWriter
    MAGIC_PDF_AVAILABLE = True
except ImportError:
    # Check for the local deployment specified by the user
    if os.path.exists(LOCAL_MINERU_PYTHON):
        MAGIC_PDF_AVAILABLE = True

def install_package(package_name: str):
    """Interactively install a missing package."""
    ans = input(f"\n[!] Missing dependency: {package_name}. Install now? (y/n): ").lower()
    if ans == 'y':
        print(f"Installing {package_name}...")
        subprocess.run([sys.executable, "-m", "pip", "install", package_name])
        return True
    return False

import psutil
import shutil

CONFIG_PATH = Path.home() / ".knowledge2skills_config.json"

def load_config():
    if CONFIG_PATH.exists():
        try:
            return json.loads(CONFIG_PATH.read_text())
        except:
            return {}
    return {}

def save_config(config):
    CONFIG_PATH.write_text(json.dumps(config, indent=2))

def check_hardware():
    """Check if RAM and Disk are sufficient for MinerU."""
    # RAM check: recommend 16GB+, minimum 8GB
    total_ram = psutil.virtual_memory().total / (1024**3)
    # Disk check: need ~10GB for models
    free_disk = shutil.disk_usage(Path.home()).free / (1024**3)
    
    return {
        "sufficient": total_ram >= 8 and free_disk >= 10,
        "ram": total_ram,
        "disk": free_disk
    }

def is_pdf_complex(path: Path) -> bool:
    """Heuristic to detect complex PDF (tables, formulas, images)."""
    if not pdfplumber:
        return False
    try:
        with pdfplumber.open(path) as pdf:
            # Check first 5 pages for complexity signals
            for i, page in enumerate(pdf.pages[:5]):
                # 1. Check for tables
                if page.extract_tables(): return True
                # 2. Check for images/rects (sign of complex layout)
                if len(page.images) > 2 or len(page.rects) > 10: return True
                # 3. Check for math symbols (crude check)
                text = page.extract_text() or ""
                if re.search(r'[∫∑√πθλαβγδεζ]', text): return True
    except:
        pass
    return False

def extract_with_mineru(path: Path, output_dir: Path) -> dict:
    """Intelligent high-precision extraction using MinerU."""
    config = load_config()
    
    # 1. If not forced by user via --high-precision, check if it's actually complex
    # (Note: we'll handle the 'high_precision' flag being passed in main)
    
    # 2. Check if already configured/used successfully
    mineru_ready = os.path.exists(LOCAL_MINERU_PYTHON) or MAGIC_PDF_AVAILABLE
    
    if not mineru_ready:
        # Check hardware
        hw = check_hardware()
        if not hw["sufficient"]:
            print(f"  [MinerU] Hardware insufficient (RAM: {hw['ram']:.1f}GB, Disk: {hw['disk']:.1f}GB). Skipping.")
            return extract_from_pdf(path)
        
        # Hardware sufficient but not installed, ask user
        print(f"\n[?] Complex PDF detected. Your hardware is sufficient for MinerU (High-Precision Parsing).")
        ans = input("Would you like to download and deploy MinerU? (This takes ~5-10GB disk space) (y/n): ").lower()
        if ans == 'y':
            print("To deploy MinerU, please follow the official guide or run the following manually:")
            print("1. git clone https://github.com/opendatalab/MinerU")
            print("2. pip install magic-pdf[full] --extra-index-url https://wheels.myhloli.com")
            print("For now, falling back to standard parser.")
            return extract_from_pdf(path)
        else:
            config["mineru_disabled"] = True
            save_config(config)
            return extract_from_pdf(path)

    # 3. Actually run MinerU
    print(f"  [High-Precision] Using MinerU to parse {path.name}...")
    
    result = None
    # Try local CLI invocation
    if os.path.exists(LOCAL_MINERU_PYTHON):
        try:
            cmd = [
                LOCAL_MINERU_PYTHON, "-m", "mineru.cli.client",
                "-p", str(path),
                "-o", str(output_dir),
                "-b", "pipeline",
                "--device", "cpu"
            ]
            run_kwargs = {"check": True, "capture_output": True}
            if MINERU_LOCAL_TIMEOUT_SECONDS > 0:
                run_kwargs["timeout"] = MINERU_LOCAL_TIMEOUT_SECONDS
            subprocess.run(cmd, **run_kwargs)
            parsed_dir = output_dir / path.stem
            md_files = list(parsed_dir.glob("*.md"))
            if md_files:
                full_md = md_files[0].read_text(encoding="utf-8")
                result = {
                    "sections": [{"heading": "MinerU Parsed Content", "content": full_md, "level": 1}],
                    "full_text": full_md,
                    "tables": [], 
                    "method": "mineru_local_cli"
                }
                # Record success
                config["mineru_last_success"] = True
                save_config(config)
        except subprocess.TimeoutExpired:
            print(
                f"  Local MinerU CLI timed out after "
                f"{MINERU_LOCAL_TIMEOUT_SECONDS}s. Falling back."
            )
        except Exception as e:
            print(f"  Local MinerU CLI failed: {e}")

    # Fallback to library import... (rest of logic same as before)
    if not result:
        try:
            from magic_pdf.pipe.UNIPipe import UNIPipe
            from magic_pdf.rw.AbsReaderWriter import FileReaderWriter
            pdf_data = path.read_bytes()
            reader = FileReaderWriter(str(output_dir))
            pipe = UNIPipe(pdf_data, {"_pdf_type": "read"}, reader)
            pipe.pipe_parse()
            content_list = pipe.pipe_mk_markdown(str(output_dir), path.name)
            full_md = ""
            for item in content_list:
                if isinstance(item, dict) and item.get('type') == 'text':
                    full_md += item['content'] + "\n\n"
                elif isinstance(item, str):
                    full_md += item + "\n\n"
            result = {
                "sections": [{"heading": "MinerU Parsed Content", "content": full_md, "level": 1}],
                "full_text": full_md,
                "tables": [], 
                "method": "mineru_library"
            }
            config["mineru_last_success"] = True
            save_config(config)
        except:
            print("  MinerU library parsing failed. Falling back.")
            return extract_from_pdf(path)
            
    return result

def extract_from_pdf(path: Path) -> dict:
    """Extract content from PDF using standard libraries."""
    sections = []
    full_text = ""
    tables = []
    
    if pdfplumber:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                full_text += text + "\n\n"
                sections.append({"heading": f"Page {i+1}", "content": text, "level": 2})
                try:
                    p_tables = page.extract_tables()
                    for t in p_tables:
                        if t: tables.append({"page": i+1, "data": t})
                except: pass
    elif PdfReader:
        reader = PdfReader(path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            full_text += text + "\n\n"
            sections.append({"heading": f"Page {i+1}", "content": text, "level": 2})
    
    return {"sections": sections, "full_text": full_text, "tables": tables}

def extract_from_epub(path: Path) -> dict:
    """Extract content from EPUB."""
    global epub
    import ebooklib
    from ebooklib import epub as epub_lib
    from bs4 import BeautifulSoup

    book = epub_lib.read_epub(str(path))
    full_text = ""
    sections = []

    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text = soup.get_text("\n")
            text = re.sub(r'\n{3,}', '\n\n', text).strip()
            if not text:
                continue
            full_text += text + "\n\n"
            heading = soup.find(['h1', 'h2'])
            title = heading.get_text().strip() if heading else f"Chapter {len(sections)+1}"
            if len(text) < 30 and title.lower().startswith("chapter"):
                continue
            sections.append({"heading": title, "content": text, "level": 2})

    return {"sections": sections, "full_text": full_text, "tables": []}


def extract_from_mobi(path: Path) -> dict:
    """Extract content from MOBI/AZW3."""
    global mobi
    if not mobi:
        if install_package("mobi"):
            import mobi
        else:
            return {"sections": [], "full_text": "Error: mobi library not installed", "tables": []}

    try:
        tmp_dir, extracted_path = mobi.extract(str(path))
        extracted = Path(extracted_path)
        if extracted.exists():
            if extracted.suffix.lower() == ".epub":
                return extract_from_epub(extracted)

            raw = extracted.read_text(encoding="utf-8", errors="ignore")
            if extracted.suffix.lower() in {".html", ".xhtml", ".htm"}:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(raw, 'html.parser')
                text = soup.get_text("\n")
            else:
                text = raw
            return {"sections": [{"heading": "Book Content", "content": text, "level": 2}], "full_text": text, "tables": []}

        return {"sections": [], "full_text": f"Error parsing: extracted file not found at {extracted_path}", "tables": []}
    except Exception as e:
        return {"sections": [], "full_text": f"Error parsing: {e}", "tables": []}


def _encode_image_for_api(path: Path, max_data_url_chars: int = 180000) -> str:
    """Resize/compress an image so it fits NVIDIA/OpenAI-compatible data URL limits."""
    mime_type, _ = mimetypes.guess_type(path.name)
    mime_type = mime_type or "image/png"
    raw = path.read_bytes()
    data_url = f"data:{mime_type};base64,{base64.b64encode(raw).decode('ascii')}"
    if len(data_url) <= max_data_url_chars:
        return data_url

    if not Image:
        raise RuntimeError("Pillow is required to resize images for analysis")

    with Image.open(path) as img:
        img = img.convert("RGB")
        quality = 70
        max_side = 1600
        while True:
            resized = img.copy()
            resized.thumbnail((max_side, max_side))
            from io import BytesIO
            buf = BytesIO()
            resized.save(buf, format="JPEG", quality=quality, optimize=True)
            payload = buf.getvalue()
            data_url = f"data:image/jpeg;base64,{base64.b64encode(payload).decode('ascii')}"
            if len(data_url) <= max_data_url_chars:
                return data_url
            if max_side <= 640 and quality <= 35:
                break
            if quality > 35:
                quality -= 10
            else:
                max_side -= 160

    raise RuntimeError(f"Image too large to analyze after compression: {path.name}")


def _clean_image_analysis_text(text: str) -> str:
    """Normalize repetitive image-analysis output from multimodal models."""
    cleaned_lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("可见地名/文字："):
            value = stripped.split("：", 1)[1]
            parts = re.split(r"[，,、；;]\s*", value)
            seen = []
            for part in parts:
                part = part.strip()
                if not part or part in seen:
                    continue
                seen.append(part)
            stripped = "可见地名/文字：" + "、".join(seen[:12])
        cleaned_lines.append(stripped)
    return "\n".join(cleaned_lines)


def extract_from_image(path: Path) -> dict:
    """Analyze standalone images such as maps so they become queryable references."""
    metadata_lines = [f"File: {path.name}"]
    if Image:
        try:
            with Image.open(path) as img:
                metadata_lines.append(f"Dimensions: {img.width}x{img.height}")
                metadata_lines.append(f"Mode: {img.mode}")
        except Exception:
            pass

    api_key = os.environ.get("OPENAI_API_KEY")
    base_url = (
        os.environ.get("OPENAI_API_BASE")
        or os.environ.get("OPENAI_BASE_URL")
        or "https://integrate.api.nvidia.com/v1"
    ).rstrip("/")
    model = os.environ.get("IMAGE_ANALYSIS_MODEL", "meta/llama-3.2-90b-vision-instruct")

    if not api_key or not requests:
        content = "\n".join(metadata_lines + ["Analysis unavailable: missing OPENAI_API_KEY or requests package."])
        return {
            "sections": [{"heading": "Image Reference", "content": content, "level": 2}],
            "full_text": content,
            "tables": [],
        }

    data_url = _encode_image_for_api(path)
    prompt = (
        "请用中文输出一个紧凑的历史地图/图片说明，严格分成这5行：\n"
        "1. 主题：\n"
        "2. 范围：\n"
        "3. 可见地名/文字：\n"
        "4. 关键信息：\n"
        "5. 与伯罗奔尼撒战争研究的关系：\n"
        "要求：不要重复地名，最多列出12个可辨认的名称，不要编造看不清的细节。"
    )
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        "temperature": 0.1,
        "max_tokens": 260,
    }

    try:
        resp = requests.post(
            f"{base_url}/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json=payload,
            timeout=120,
        )
        resp.raise_for_status()
        content = _clean_image_analysis_text(resp.json()["choices"][0]["message"]["content"].strip())
    except Exception as e:
        content = "\n".join(metadata_lines + [f"Analysis failed: {e}"])
    else:
        content = "\n".join(metadata_lines + ["", content])

    return {
        "sections": [{"heading": "Image Reference", "content": content, "level": 2}],
        "full_text": content,
        "tables": [],
    }

def extract_from_docx(path: Path) -> dict:
    """Extract content from DOCX."""
    if not docx: return {"sections": [], "full_text": "Error: python-docx not installed", "tables": []}
    doc = docx.Document(path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    sections = []
    curr_sec = {"heading": "Document", "content": [], "level": 1}
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if curr_sec["content"]:
                curr_sec["content"] = "\n".join(curr_sec["content"])
                sections.append(curr_sec)
            curr_sec = {"heading": para.text, "content": [], "level": 2}
        else:
            curr_sec["content"].append(para.text)
    
    if curr_sec["content"]:
        curr_sec["content"] = "\n".join(curr_sec["content"])
        sections.append(curr_sec)
        
    return {"sections": sections, "full_text": full_text, "tables": []}

def extract_from_text(path: Path) -> dict:
    """Extract from TXT or MD."""
    text = path.read_text(encoding="utf-8", errors="ignore")
    lines = text.split("\n")
    sections = []
    curr_sec = {"heading": "Start", "content": [], "level": 2}
    
    for line in lines:
        if line.startswith("#"):
            if curr_sec["content"]:
                curr_sec["content"] = "\n".join(curr_sec["content"])
                sections.append(curr_sec)
            curr_sec = {"heading": line.lstrip("# ").strip(), "content": [], "level": line.count("#")}
        else:
            curr_sec["content"].append(line)
            
    if curr_sec["content"]:
        curr_sec["content"] = "\n".join(curr_sec["content"])
        sections.append(curr_sec)
        
    return {"sections": sections, "full_text": text, "tables": []}


def _coerce_text_chunks(value) -> list[str]:
    """Collect readable text fragments from nested JSON values."""
    if value is None:
        return []
    if isinstance(value, str):
        stripped = value.strip()
        return [stripped] if stripped else []
    if isinstance(value, (int, float, bool)):
        return [str(value)]
    if isinstance(value, list):
        chunks = []
        for item in value:
            chunks.extend(_coerce_text_chunks(item))
        return chunks
    if isinstance(value, dict):
        chunks = []
        for key in [
            "text",
            "content",
            "body",
            "caption",
            "table_caption",
            "img_caption",
            "image_caption",
            "latex",
            "html",
            "markdown",
            "md",
            "footnote",
            "ocr_text",
        ]:
            if key in value:
                chunks.extend(_coerce_text_chunks(value.get(key)))
        return chunks
    return []


def _normalize_block_type(raw_type: str) -> str:
    value = (raw_type or "block").strip().lower().replace("-", "_")
    alias_map = {
        "image": "image",
        "img": "image",
        "figure": "image",
        "picture": "image",
        "table": "table",
        "form": "table",
        "equation": "equation",
        "formula": "equation",
        "latex": "equation",
        "text": "text",
        "paragraph": "text",
        "title": "heading",
        "heading": "heading",
        "header": "heading",
        "list": "list",
    }
    return alias_map.get(value, value)


def _find_structured_items(payload):
    """Return a likely content-list array from MinerU/RAG-style JSON."""
    if isinstance(payload, list):
        return payload
    if not isinstance(payload, dict):
        return None

    for key in ["content_list_v2", "content_list", "items", "data", "blocks"]:
        value = payload.get(key)
        if isinstance(value, list):
            return value

    for value in payload.values():
        if isinstance(value, list) and value and all(isinstance(item, dict) for item in value):
            return value

    return None


def _build_section_path(item: dict) -> str:
    path_keys = [
        "section_path",
        "section_title",
        "chapter_title",
        "parent_title",
        "doc_title",
    ]
    path_parts = []
    for key in path_keys:
        value = item.get(key)
        if isinstance(value, list):
            path_parts.extend(_coerce_text_chunks(value))
        else:
            path_parts.extend(_coerce_text_chunks(value))
    deduped = []
    for part in path_parts:
        if part not in deduped:
            deduped.append(part)
    return " > ".join(deduped[:6])


def _render_structured_item(item: dict, block_index: int) -> tuple[str, str]:
    """Render a content-list item into a heading/content pair."""
    block_type = _normalize_block_type(
        item.get("type") or item.get("category") or item.get("block_type") or "block"
    )
    page = item.get("page_idx") or item.get("page") or item.get("page_no") or item.get("page_num")
    section_path = _build_section_path(item)

    candidate_text = []
    for key in [
        "text",
        "content",
        "body",
        "md",
        "markdown",
        "latex",
        "caption",
        "table_caption",
        "img_caption",
        "image_caption",
        "html",
        "footnote",
        "ocr_text",
        "table_body",
        "table_markdown",
        "table_html",
    ]:
        candidate_text.extend(_coerce_text_chunks(item.get(key)))

    if not candidate_text and "table" in block_type:
        rows = item.get("table") or item.get("data") or item.get("cells")
        candidate_text.extend(_coerce_text_chunks(rows))

    if not candidate_text:
        return "", ""

    content_lines = [f"Block Type: {block_type}"]
    if page not in [None, ""]:
        content_lines.append(f"Page: {page}")
    if section_path:
        content_lines.append(f"Section Path: {section_path}")

    content_lines.append("")
    content_lines.extend(candidate_text)

    title_text = ""
    if block_type == "heading":
        title_text = candidate_text[0][:80]
    elif section_path:
        title_text = section_path.split(" > ")[-1][:80]

    if not title_text:
        title_text = f"{block_type.replace('_', ' ').title()} Block {block_index + 1}"

    content = "\n".join(line for line in content_lines if line is not None).strip()
    return title_text, content


def extract_from_json(path: Path) -> dict:
    """Extract from raw JSON or structured content-list JSON."""
    try:
        raw_data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {"sections": [], "full_text": "", "tables": []}

    structured_items = _find_structured_items(raw_data)
    if structured_items:
        sections = []
        full_text_parts = []
        tables = []

        for index, item in enumerate(structured_items):
            if not isinstance(item, dict):
                continue
            heading, content = _render_structured_item(item, index)
            if not content:
                continue
            block_type = _normalize_block_type(
                item.get("type") or item.get("category") or item.get("block_type") or "block"
            )
            sections.append(
                {
                    "heading": heading,
                    "content": content,
                    "level": 2,
                    "block_type": block_type,
                }
            )
            full_text_parts.append(content)
            if block_type == "table":
                table_payload = (
                    item.get("table")
                    or item.get("data")
                    or item.get("cells")
                    or item.get("table_body")
                    or item.get("table_markdown")
                    or item.get("table_html")
                )
                tables.append(
                    {
                        "page": item.get("page_idx") or item.get("page") or item.get("page_no"),
                        "data": table_payload,
                    }
                )

        if sections:
            return {
                "sections": sections,
                "full_text": "\n\n".join(full_text_parts),
                "tables": tables,
            }

    if isinstance(raw_data, dict):
        summary_lines = [f"{key}: {json.dumps(value, ensure_ascii=False)}" for key, value in raw_data.items()]
        summary_text = "\n".join(summary_lines)
        return {
            "sections": [{"heading": "JSON Content", "content": summary_text, "level": 2}],
            "full_text": summary_text,
            "tables": [],
        }

    text = json.dumps(raw_data, ensure_ascii=False, indent=2)
    return {
        "sections": [{"heading": "JSON Content", "content": text, "level": 2}],
        "full_text": text,
        "tables": [],
    }

def process_files(paths: list, high_precision: bool = False, work_dir: Path = None) -> dict:
    """Process multiple files into a single unified extraction."""
    combined = {
        "metadata": {"sources": [], "total_files": len(paths)},
        "sections": [],
        "tables": [],
        "full_text": ""
    }
    
    for p_str in paths:
        path = Path(p_str).resolve()
        if not path.exists(): continue
        
        ext = path.suffix.lower()
        combined["metadata"]["sources"].append({"name": path.name, "type": ext})
        
        print(f"Processing: {path.name}...")
        
        if ext == ".pdf":
            # Auto-detect complexity
            complex_pdf = is_pdf_complex(path)
            config = load_config()
            
            # Use MinerU if: 
            # 1. User forced it via --high-precision
            # 2. OR it's complex AND (it worked before OR user didn't disable it)
            should_use_mineru = high_precision or (
                complex_pdf and 
                not config.get("mineru_disabled")
            )
            
            if should_use_mineru and work_dir:
                res = extract_with_mineru(path, work_dir)
            else:
                if complex_pdf:
                    print(f"  [Notice] {path.name} appears complex, but MinerU is skipped.")
                res = extract_from_pdf(path)
        elif ext == ".docx": res = extract_from_docx(path)
        elif ext == ".epub": res = extract_from_epub(path)
        elif ext in [".mobi", ".azw3"]: res = extract_from_mobi(path)
        elif ext in [".png", ".jpg", ".jpeg", ".webp"]: res = extract_from_image(path)
        elif ext in [".txt", ".md", ".markdown"]: res = extract_from_text(path)
        elif ext == ".json": res = extract_from_json(path)
        else: continue
        
        for s in res.get("sections", []):
            s["heading"] = f"[{path.name}] {s['heading']}"
            combined["sections"].append(s)
        
        combined["full_text"] += f"\n\n--- Source: {path.name} ---\n\n" + res.get("full_text", "")
        combined["tables"].extend(res.get("tables", []))
        
    return combined

def main():
    parser = argparse.ArgumentParser(description="Extract content from various formats")
    parser.add_argument("files", nargs="+", help="Paths to files to process")
    parser.add_argument("--output", "-o", help="Output JSON path")
    parser.add_argument("--high-precision", action="store_true", help="Use MinerU for PDFs")
    args = parser.parse_args()
    
    # Need a work dir for MinerU
    work_dir = Path("./extraction_work").resolve()
    work_dir.mkdir(exist_ok=True)
    
    result = process_files(args.files, high_precision=args.high_precision, work_dir=work_dir)
    
    if args.output:
        Path(args.output).write_text(json.dumps(result, ensure_ascii=False, indent=2))
        print(f"Saved to {args.output}")
    else:
        print(json.dumps(result, indent=2))

if __name__ == "__main__":
    main()
